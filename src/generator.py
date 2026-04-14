"""Document generation layer for JobRadar.

Generates tailored cover letters and CV changes using the AI client.
All terminal interaction lives in main.py — this module is a library.
No print() statements except the heading-not-found warning in apply_cv_changes().
"""

import re
from datetime import datetime
from pathlib import Path

from docx import Document

from src.ai_client import complete, complete_json
from src.db import get_example_letters


def _slugify(value: str, max_len: int = 30) -> str:
    """Lowercase, replace non-alphanumeric chars with hyphens, truncate."""
    value = value.lower()
    value = re.sub(r'[^a-z0-9]+', '-', value)
    value = value.strip('-')
    return value[:max_len].rstrip('-')


BANNED_PHRASES = [
    'leverage', 'leveraging', 'excited', 'passionate', 'thrilled',
    'synergy', 'innovative', 'dynamic', 'results-driven', 'detail-oriented',
    'team player', 'hard worker', 'proven track record', 'fast-paced environment',
    'i am writing to express', 'i would like to apply', 'i am confident that',
    'i believe i would be a great fit', 'stabilizing force',
]


def build_cover_letter_prompt(
    profile: dict,
    jd_text: str,
    example_letters: list[dict],
    language: str = 'en',
) -> str:
    """Assemble the cover letter generation prompt.

    Reads config/prompts/generate.txt at runtime — never cached.
    Substitutes all {placeholders} with profile data and examples.

    profile['voice'] is the top-level 'voice:' key from profile.yaml.
    It is a required key — do not use .get() with a fallback.

    example_letters: list of dicts from db.get_example_letters().
    Each dict has keys: id, job_id, doc_type, path, version,
                        rating, use_as_example, created_at.
    """
    prompt_path = Path('config/prompts/generate.txt')
    if not prompt_path.exists():
        raise FileNotFoundError("config/prompts/generate.txt not found.")
    template = prompt_path.read_text(encoding='utf-8')

    # Build examples_section
    if not example_letters:
        examples_section = "No rated examples yet — use the voice document as the only style guide."
    else:
        blocks = []
        for letter in example_letters:
            path = letter['path']
            ext = Path(path).suffix.lower()
            if ext == '.pdf':
                continue
            try:
                if ext == '.docx':
                    doc = Document(path)
                    text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
                else:
                    text = Path(path).read_text(encoding='utf-8')
                blocks.append(
                    f"--- Example letter (rating {letter['rating']}/5) ---\n{text}\n"
                )
            except Exception:
                continue
        if blocks:
            examples_section = '\n\n'.join(blocks)
        else:
            examples_section = "No rated examples yet — use the voice document as the only style guide."

    return template.format(
        name=profile['personal']['name'],
        age=profile['personal']['age'],
        location=profile['personal']['location'],
        voice=profile['voice'],
        skills_expert=', '.join(profile['skills']['expert']),
        skills_solid=', '.join(profile['skills']['solid']),
        job_text=jd_text,
        language='German' if language == 'de' else 'English',
        examples_section=examples_section,
    )


def build_cv_prompt(
    profile: dict,
    jd_text: str,
    base_cv_text: str,
) -> str:
    """Assemble the CV tailoring prompt.

    Reads config/prompts/cv.txt at runtime — never cached.
    Substitutes all {placeholders} with profile data.
    The AI must respond with JSON only — parsed by the caller via complete_json().
    """
    prompt_path = Path('config/prompts/cv.txt')
    if not prompt_path.exists():
        raise FileNotFoundError("config/prompts/cv.txt not found.")
    template = prompt_path.read_text(encoding='utf-8')

    return template.format(
        name=profile['personal']['name'],
        skills_expert=', '.join(profile['skills']['expert']),
        skills_solid=', '.join(profile['skills']['solid']),
        voice=profile['voice'],
        base_cv_text=base_cv_text,
        job_text=jd_text,
    )


def generate_cover_letter(
    client,
    profile: dict,
    jd_text: str,
    language: str = 'en',
) -> str:
    """Generate a cover letter using the AI.

    Fetches up to 3 highest-rated example letters from DB via get_example_letters().
    Returns the cover letter as a plain text string.
    Raises ValueError if the response contains a banned phrase.
    """
    example_letters = get_example_letters(min_rating=4, limit=3)
    prompt = build_cover_letter_prompt(profile, jd_text, example_letters, language)
    text = complete(client, prompt)

    for phrase in BANNED_PHRASES:
        if re.search(r'\b' + re.escape(phrase) + r'\b', text, re.IGNORECASE):
            raise ValueError(
                f"Banned phrase found in output: '{phrase}'. Regenerate or edit manually."
            )

    return text


def generate_cv_changes(
    client,
    profile: dict,
    jd_text: str,
) -> dict:
    """Ask the AI to propose targeted CV changes for this JD.

    Returns a dict with exactly these keys:
      profile_summary, skills_to_highlight, skills_to_remove, changes_explained

    Internally calls extract_base_cv_text(Path('examples/cv_base.md')) to
    supply base_cv_text to build_cv_prompt(). If cv_base.md is missing,
    FileNotFoundError propagates to the caller — do NOT catch it here.

    Raises ValueError if:
      - complete_json() cannot parse the response (propagated unchanged)
      - the returned dict is missing any required key

    Does NOT catch any exceptions — let them propagate to the caller (main.py Step 4).
    """
    base_cv_text = extract_base_cv_text(Path('examples/cv_base.md'))
    prompt = build_cv_prompt(profile, jd_text, base_cv_text)
    result = complete_json(client, prompt)

    REQUIRED_CV_KEYS = {'profile_summary', 'skills_to_highlight', 'skills_to_remove', 'changes_explained'}
    missing = REQUIRED_CV_KEYS - result.keys()
    if missing:
        raise ValueError(f"AI response missing required keys: {missing}")

    return result


def extract_base_cv_text(cv_path: Path) -> str:
    """Read plain text from examples/cv_base.md for use in the CV prompt.

    Raises FileNotFoundError if the file does not exist.
    Never modifies the source file.
    """
    if not cv_path.exists():
        raise FileNotFoundError(f"CV base not found: {cv_path}")
    return cv_path.read_text(encoding='utf-8')


def apply_cv_changes(
    cv_path: Path,
    changes: dict,
    output_path: Path,
) -> None:
    """Apply the AI-proposed profile summary to the base CV and save as new .md.

    Replaces the paragraph immediately after '## Professional Profile'.
    All other content preserved exactly.
    Profile summary must be a single unwrapped line in cv_base.md.
    Reads from cv_path. Writes to output_path. Never touches cv_base.md.
    """
    lines = cv_path.read_text(encoding='utf-8').splitlines()
    output_lines: list[str] = []
    in_profile_section = False
    replaced = False
    skip_until_next_content = False

    for line in lines:
        stripped = line.strip()

        # Detect the profile section heading
        if stripped.lower() in ('## professional profile', '## profil'):
            in_profile_section = True
            skip_until_next_content = True
            output_lines.append(line)
            continue

        # First non-empty, non-heading line after the profile heading
        if in_profile_section and skip_until_next_content:
            if stripped and not stripped.startswith('#'):
                output_lines.append(changes['profile_summary'])
                skip_until_next_content = False
                in_profile_section = False
                replaced = True
                continue  # skip the old summary line
            elif stripped.startswith('#'):
                # Heading immediately after — no body paragraph found
                skip_until_next_content = False
                in_profile_section = False

        output_lines.append(line)

    if not replaced:
        output_lines.append('')
        output_lines.append(changes['profile_summary'])
        print(
            "Warning: '## Professional Profile' heading not found "
            "— summary appended. Review before sending."
        )

    output_path.write_text('\n'.join(output_lines), encoding='utf-8')


def write_cover_letter_docx(
    text: str,
    output_path: Path,
    profile: dict,
    company: str,
    date_str: str,
) -> None:
    """Write cover letter text to a .docx file.

    Structure: sender block, date, company, salutation, body, sign-off.
    date_str: YYYYMMDD format.
    """
    personal = profile['personal']
    name = personal['name']
    date_display = datetime.strptime(date_str, "%Y%m%d").strftime("%B %Y")
    # Build day without leading zero (cross-platform)
    day = str(datetime.strptime(date_str, "%Y%m%d").day)
    month_year = datetime.strptime(date_str, "%Y%m%d").strftime("%B %Y")
    date_display = f"{day} {month_year}"

    doc = Document()

    # Sender block
    doc.add_paragraph(name)
    doc.add_paragraph(f"{personal['email']}  ·  {personal['phone']}")
    doc.add_paragraph(personal['location'])
    doc.add_paragraph('')

    # Date and recipient
    doc.add_paragraph(date_display)
    doc.add_paragraph(company)
    doc.add_paragraph('')

    # Salutation
    doc.add_paragraph(f"Dear {company} Team,")
    doc.add_paragraph('')

    # Body
    for chunk in text.split('\n\n'):
        if chunk.strip():
            doc.add_paragraph(chunk.strip())

    # Sign-off
    doc.add_paragraph('')
    doc.add_paragraph('Best regards,')
    doc.add_paragraph('')
    doc.add_paragraph(name)

    doc.save(output_path)


def write_cover_letter_md(
    text: str,
    output_path: Path,
    body_path: Path,
    profile: dict,
    company: str,
    date_str: str,
    language: str = 'en',
) -> None:
    """Write cover letter to two .md files.

    output_path — full letter with sender block and framing. Used for PDF.
    body_path   — body paragraphs only. Used as style example in DB.

    date_str: YYYYMMDD format.
    Uses explicit <!-- SENDER_START/END --> delimiters so the PDF converter
    never needs to guess structure.

    Coupled to cover_letter_md_to_pdf() in src/pdf_writer.py — any structural
    change here must be reflected there in the same commit.
    """
    personal = profile['personal']
    name = personal['name']
    day = str(datetime.strptime(date_str, "%Y%m%d").day)
    month_year = datetime.strptime(date_str, "%Y%m%d").strftime("%B %Y")
    date_display = f"{day} {month_year}"

    if language == 'de':
        salutation = "Sehr geehrte Damen und Herren,"
        signoff = "Mit freundlichen Grüßen,"
    else:
        salutation = f"Dear {company} Team,"
        signoff = "Best regards,"

    # Normalise body — split on double newlines, strip each chunk
    body_paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    body_text = '\n\n'.join(body_paragraphs)

    # Full letter with explicit delimiters
    full_lines = [
        "<!-- SENDER_START -->",
        f"**{name}**",
        f"{personal['email']} \u00b7 {personal['phone']}",
        f"{personal['location']}",
        "<!-- SENDER_END -->",
        "",
        date_display,
        "",
        company,
        "",
        salutation,
        "",
        body_text,
        "",
        signoff,
        "",
        name,
    ]
    output_path.write_text('\n'.join(full_lines), encoding='utf-8')

    # Body only — for example pool (no framing boilerplate)
    body_path.write_text(body_text, encoding='utf-8')


def make_output_dir(company: str, role_title: str, date_str: str) -> Path:
    """Build and return the output directory Path. Does NOT create it.

    Caller is responsible for mkdir().

    Naming: output/{date_str}/{company-slug}/
    date_str format: YYYYMMDD
    """
    return Path('output') / date_str / _slugify(company)
